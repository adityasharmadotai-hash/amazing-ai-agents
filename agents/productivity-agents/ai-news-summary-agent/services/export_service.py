"""Export service — render content to PDF, DOCX, or Markdown bytes.

All methods return `bytes` so they can be streamed directly into a Streamlit
`st.download_button` without touching disk.
"""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from utils.logging_config import get_logger

logger = get_logger(__name__)


def _md_to_lines(markdown_text: str) -> list[tuple[str, str]]:
    """Very small Markdown tokenizer → list of (kind, text)."""
    lines: list[tuple[str, str]] = []
    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            lines.append(("space", ""))
        elif line.startswith("### "):
            lines.append(("h3", line[4:]))
        elif line.startswith("## "):
            lines.append(("h2", line[3:]))
        elif line.startswith("# "):
            lines.append(("h1", line[2:]))
        elif re.match(r"^\s*[-*]\s+", line):
            lines.append(("bullet", re.sub(r"^\s*[-*]\s+", "", line)))
        else:
            lines.append(("p", line))
    return lines


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


class ExportService:
    # ---------------------------- Markdown ---------------------------
    @staticmethod
    def to_markdown(title: str, content: str) -> bytes:
        stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        doc = f"# {title}\n\n_Generated {stamp}_\n\n{content}\n"
        return doc.encode("utf-8")

    # ------------------------------ DOCX -----------------------------
    @staticmethod
    def to_docx(title: str, content: str) -> bytes:
        document = Document()
        heading = document.add_heading(title, level=0)
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        meta = document.add_paragraph(
            f"Generated {datetime.now().strftime('%d %b %Y, %H:%M')}"
        )
        meta.runs[0].italic = True
        meta.runs[0].font.size = Pt(9)

        for kind, text in _md_to_lines(content):
            text = _strip_inline_md(text)
            if kind == "h1":
                document.add_heading(text, level=1)
            elif kind == "h2":
                document.add_heading(text, level=2)
            elif kind == "h3":
                document.add_heading(text, level=3)
            elif kind == "bullet":
                document.add_paragraph(text, style="List Bullet")
            elif kind == "space":
                continue
            else:
                document.add_paragraph(text)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    # ------------------------------ PDF ------------------------------
    @staticmethod
    def to_pdf(title: str, content: str) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=2 * cm, bottomMargin=2 * cm,
            leftMargin=2 * cm, rightMargin=2 * cm,
            title=title,
        )
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle("EmailBody", parent=styles["Normal"], fontSize=10,
                                  leading=15, alignment=TA_LEFT))
        styles.add(ParagraphStyle("EmailBullet", parent=styles["Normal"], fontSize=10,
                                  leading=15, leftIndent=14, bulletIndent=4))

        flow = [Paragraph(title, styles["Title"]),
                Paragraph(datetime.now().strftime("%d %b %Y, %H:%M"), styles["Italic"]),
                Spacer(1, 0.4 * cm)]

        for kind, text in _md_to_lines(content):
            text = _strip_inline_md(text).replace("&", "&amp;").replace("<", "&lt;")
            if kind == "h1":
                flow.append(Paragraph(text, styles["Heading1"]))
            elif kind == "h2":
                flow.append(Paragraph(text, styles["Heading2"]))
            elif kind == "h3":
                flow.append(Paragraph(text, styles["Heading3"]))
            elif kind == "bullet":
                flow.append(Paragraph(f"• {text}", styles["EmailBullet"]))
            elif kind == "space":
                flow.append(Spacer(1, 0.2 * cm))
            else:
                flow.append(Paragraph(text, styles["EmailBody"]))

        doc.build(flow)
        return buffer.getvalue()
