"""
exporter.py — Export proposals to branded PDF and editable DOCX.
"""

import io
from datetime import datetime


def _date() -> str:
    return datetime.now().strftime("%B %d, %Y")


def _hex_to_rgb(h: str) -> tuple:
    h = h.lstrip("#")
    return tuple(int(h[i:i+2],16) for i in (0,2,4))


def _safe(text) -> str:
    if not text: return ""
    return str(text).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")


# ── PDF ────────────────────────────────────────────────────────────────────────

def export_pdf(proposal: dict, branding: dict) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
        )

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()

        pr  = _hex_to_rgb(branding.get("primary_color","#4f46e5"))
        sr  = _hex_to_rgb(branding.get("secondary_color","#7c3aed"))
        PRI = colors.Color(pr[0]/255, pr[1]/255, pr[2]/255)
        SEC = colors.Color(sr[0]/255, sr[1]/255, sr[2]/255)
        DARK  = colors.HexColor("#0f172a")
        GRAY  = colors.HexColor("#64748b")
        LIGHT = colors.HexColor("#f8fafc")
        WHITE = colors.white

        def sty(name, **kw):
            return ParagraphStyle(name, parent=styles["Normal"], **kw)

        H1   = sty("H1",  fontSize=22, fontName="Helvetica-Bold", textColor=DARK, spaceAfter=4, leading=28)
        H2   = sty("H2",  fontSize=15, fontName="Helvetica-Bold", textColor=PRI, spaceAfter=6, spaceBefore=14, leading=20)
        H3   = sty("H3",  fontSize=12, fontName="Helvetica-Bold", textColor=DARK, spaceAfter=4, spaceBefore=8)
        BODY = sty("Body", fontSize=10, leading=16, spaceAfter=5, textColor=DARK)
        META = sty("Meta", fontSize=10, textColor=GRAY, spaceAfter=3)
        BULL = sty("Bull", fontSize=10, leftIndent=14, spaceAfter=3, leading=14, textColor=DARK)
        SMALL= sty("Small",fontSize=9, textColor=GRAY, spaceAfter=2)
        ITAL = sty("Ital", fontSize=12, textColor=SEC, spaceAfter=6, leading=18)

        meta = proposal.get("_meta",{})
        agency = branding.get("agency_name","Your Agency")
        story  = []

        # ── Cover ──
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(_safe(agency), sty("Ag", fontSize=15, fontName="Helvetica-Bold", textColor=PRI, spaceAfter=3)))
        story.append(Paragraph(_safe(branding.get("tagline","")), META))
        story.append(HRFlowable(width="100%", thickness=3, color=PRI))
        story.append(Spacer(1, 0.25*inch))
        story.append(Paragraph(_safe(proposal.get("proposal_title","Business Proposal")), H1))
        story.append(Paragraph(_safe(proposal.get("tagline","")), ITAL))
        story.append(Spacer(1, 0.15*inch))

        # Cover table
        cdata = [
            ["Prepared for:", meta.get("client_name","")],
            ["Prepared by:",  agency],
            ["Date:",         meta.get("today",_date())],
            ["Project:",      meta.get("project_type","")],
            ["Budget:",       meta.get("budget","")],
            ["Timeline:",     meta.get("timeline","")],
        ]
        ct = Table(cdata, colWidths=[1.7*inch, 4.3*inch])
        ct.setStyle(TableStyle([
            ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"), ("FONTSIZE",(0,0),(-1,-1),10),
            ("TEXTCOLOR",(0,0),(0,-1),PRI), ("TEXTCOLOR",(1,0),(1,-1),DARK),
            ("ROWBACKGROUNDS",(0,0),(-1,-1),[LIGHT,WHITE]),
            ("TOPPADDING",(0,0),(-1,-1),7), ("BOTTOMPADDING",(0,0),(-1,-1),7),
            ("LEFTPADDING",(0,0),(-1,-1),10),
            ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#e2e8f0")),
        ]))
        story.append(ct)
        story.append(Spacer(1,0.4*inch))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")))

        def section(title, content):
            story.append(Paragraph(title, H2))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
            story.append(Spacer(1,0.04*inch))
            if isinstance(content, str):
                for chunk in content.split("\n\n"):
                    if chunk.strip():
                        story.append(Paragraph(_safe(chunk.strip()), BODY))
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str):
                        story.append(Paragraph(f"• {_safe(item)}", BULL))

        story.append(PageBreak())
        section("Executive Summary", proposal.get("executive_summary",""))
        section("About Us", proposal.get("about_us",""))
        section("The Challenge", proposal.get("problem_statement",""))
        section("Our Solution", proposal.get("proposed_solution",""))

        # Scope
        story.append(Paragraph("Scope of Work", H2))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
        for i, phase in enumerate(proposal.get("scope_of_work",[])):
            story.append(Paragraph(f"{_safe(phase.get('phase',''))}  ·  {_safe(phase.get('duration',''))}", H3))
            if phase.get("description"):
                story.append(Paragraph(_safe(phase["description"]), BODY))
            for d in phase.get("deliverables",[]):
                story.append(Paragraph(f"  ✓  {_safe(d)}", BULL))
            if i < len(proposal.get("scope_of_work",[]))-1:
                story.append(Spacer(1,0.08*inch))

        # Timeline
        story.append(PageBreak())
        story.append(Paragraph("Project Timeline", H2))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
        tl = proposal.get("timeline",[])
        if tl:
            tdata = [["Milestone","Week","Description"]]
            for m in tl:
                tdata.append([m.get("milestone",""), m.get("week",""), m.get("description","")])
            tt = Table(tdata, colWidths=[2.2*inch, 1.0*inch, 3.4*inch])
            tt.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0),PRI), ("TEXTCOLOR",(0,0),(-1,0),WHITE),
                ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"), ("FONTSIZE",(0,0),(-1,-1),9),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT,WHITE]),
                ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#e2e8f0")),
                ("TOPPADDING",(0,0),(-1,-1),6), ("BOTTOMPADDING",(0,0),(-1,-1),6),
                ("LEFTPADDING",(0,0),(-1,-1),8), ("VALIGN",(0,0),(-1,-1),"TOP"),
            ]))
            story.append(tt)

        # Pricing
        story.append(PageBreak())
        story.append(Paragraph("Investment & Pricing", H2))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
        pricing = proposal.get("pricing",{})
        if pricing.get("intro"):
            story.append(Paragraph(_safe(pricing["intro"]), BODY))
        packages = pricing.get("packages",[])
        if packages:
            pdata = [["Package","Investment","Includes"]]
            for pkg in packages:
                feats = "\n".join(f"• {f}" for f in pkg.get("features",[]))
                rec   = " ★ RECOMMENDED" if pkg.get("recommended") else ""
                pdata.append([f"{pkg.get('name','')}{rec}", pkg.get("price",""), feats])
            pt = Table(pdata, colWidths=[1.8*inch,1.4*inch,3.4*inch])
            pt.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0),SEC), ("TEXTCOLOR",(0,0),(-1,0),WHITE),
                ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"), ("FONTSIZE",(0,0),(-1,-1),9),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT,WHITE]),
                ("GRID",(0,0),(-1,-1),0.5,colors.HexColor("#e2e8f0")),
                ("TOPPADDING",(0,0),(-1,-1),8), ("BOTTOMPADDING",(0,0),(-1,-1),8),
                ("LEFTPADDING",(0,0),(-1,-1),8), ("VALIGN",(0,0),(-1,-1),"TOP"),
            ]))
            story.append(pt)
        if pricing.get("payment_terms"):
            story.append(Spacer(1,0.08*inch))
            story.append(Paragraph(f"<b>Payment Terms:</b> {_safe(pricing['payment_terms'])}", BODY))

        # Why us
        why = proposal.get("why_us",[])
        if why:
            section("Why Choose Us", [f"{r.get('title','')}: {r.get('description','')}" for r in why])

        # Case study
        cs = proposal.get("case_study",{})
        if cs and cs.get("title"):
            story.append(Paragraph("Case Study", H2))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
            story.append(Paragraph(_safe(cs["title"]), H3))
            for lbl, key in [("Challenge","challenge"),("Solution","solution"),("Result","result")]:
                if cs.get(key):
                    story.append(Paragraph(f"<b>{lbl}:</b> {_safe(cs[key])}", BODY))

        # Terms
        story.append(PageBreak())
        terms = proposal.get("terms",{})
        if terms:
            story.append(Paragraph("Terms & Conditions", H2))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
            for k, v in terms.items():
                if v:
                    story.append(Paragraph(k.replace("_"," ").title(), H3))
                    story.append(Paragraph(_safe(str(v)), BODY))

        # Next steps
        ns = proposal.get("next_steps",[])
        if ns:
            story.append(Paragraph("Next Steps", H2))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e8edf5")))
            for i, step in enumerate(ns,1):
                story.append(Paragraph(f"{i}.  {_safe(step)}", BULL))

        # Closing
        closing = proposal.get("closing_statement","")
        if closing:
            story.append(Spacer(1,0.2*inch))
            story.append(HRFlowable(width="100%", thickness=2, color=PRI))
            story.append(Paragraph(_safe(closing), sty("Cl", fontSize=11, leading=18, textColor=DARK, spaceAfter=10, spaceBefore=10)))

        contacts = [x for x in [branding.get("email",""), branding.get("phone",""), branding.get("website","")] if x]
        if contacts:
            story.append(Paragraph(" · ".join(contacts), SMALL))
        story.append(Paragraph(f"© {datetime.now().year} {_safe(agency)} · Confidential", SMALL))

        doc.build(story)
        buf.seek(0)
        return buf.read()

    except ImportError:
        return b"reportlab not installed. Run: pip install reportlab"
    except Exception as e:
        return f"PDF error: {e}".encode()


# ── DOCX ───────────────────────────────────────────────────────────────────────

def export_docx(proposal: dict, branding: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(0.75)
            sec.left_margin = sec.right_margin = Inches(1.0)

        pr  = _hex_to_rgb(branding.get("primary_color","#4f46e5"))
        sr  = _hex_to_rgb(branding.get("secondary_color","#7c3aed"))
        PRI = RGBColor(*pr); SEC = RGBColor(*sr)
        DARK= RGBColor(0x0f,0x17,0x2a); GRAY=RGBColor(0x64,0x74,0x8b)

        meta   = proposal.get("_meta",{})
        agency = branding.get("agency_name","Your Agency")

        def heading(text, level=1, color=None):
            p = doc.add_heading(str(text), level=level)
            c = color or (PRI if level<=1 else DARK)
            for r in p.runs: r.font.color.rgb = c

        def para(text, bold=False, italic=False, color=None, size=10):
            p = doc.add_paragraph()
            r = p.add_run(str(text))
            r.bold=bold; r.italic=italic; r.font.size=Pt(size)
            if color: r.font.color.rgb = color

        def bullet(text):
            doc.add_paragraph(str(text), style="List Bullet")

        heading(agency, 0); doc.add_paragraph(branding.get("tagline",""))
        doc.add_paragraph(f"Prepared for: {meta.get('client_name','')} | Date: {meta.get('today',_date())}")
        doc.add_paragraph("─"*60)
        heading(proposal.get("proposal_title","Business Proposal"), 1, PRI)
        para(proposal.get("tagline",""), italic=True, color=SEC, size=12)

        def section(title, content):
            heading(title, 2)
            doc.add_paragraph("─"*40)
            if isinstance(content, str):
                for chunk in content.split("\n\n"):
                    if chunk.strip(): para(chunk.strip())
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str): bullet(item)

        section("Executive Summary", proposal.get("executive_summary",""))
        section("About Us", proposal.get("about_us",""))
        section("The Challenge", proposal.get("problem_statement",""))
        section("Our Solution", proposal.get("proposed_solution",""))

        heading("Scope of Work", 2); doc.add_paragraph("─"*40)
        for phase in proposal.get("scope_of_work",[]):
            heading(f"{phase.get('phase','')} — {phase.get('duration','')}", 3)
            if phase.get("description"): para(phase["description"])
            for d in phase.get("deliverables",[]): bullet(f"✓ {d}")

        doc.add_page_break()
        heading("Project Timeline", 2)
        tl = proposal.get("timeline",[])
        if tl:
            table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(["Milestone","Week","Description"]):
                hdr[i].text = h
                for r in hdr[i].paragraphs:
                    for run in r.runs: run.bold=True
            for m in tl:
                row = table.add_row().cells
                row[0].text = m.get("milestone","")
                row[1].text = m.get("week","")
                row[2].text = m.get("description","")

        doc.add_page_break()
        heading("Investment & Pricing", 2)
        pricing = proposal.get("pricing",{})
        if pricing.get("intro"): para(pricing["intro"])
        packages = pricing.get("packages",[])
        if packages:
            table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(["Package","Investment","Features"]):
                hdr[i].text = h
                for r in hdr[i].paragraphs:
                    for run in r.runs: run.bold=True
            for pkg in packages:
                row = table.add_row().cells
                row[0].text = pkg.get("name","") + (" ★ Recommended" if pkg.get("recommended") else "")
                row[1].text = pkg.get("price","")
                row[2].text = "\n".join(f"• {f}" for f in pkg.get("features",[]))
        if pricing.get("payment_terms"):
            para(f"Payment Terms: {pricing['payment_terms']}", bold=True)

        why = proposal.get("why_us",[])
        if why:
            heading("Why Choose Us", 2)
            for r in why:
                p = doc.add_paragraph()
                p.add_run(f"{r.get('title','')}: ").bold=True
                p.add_run(r.get("description",""))

        cs = proposal.get("case_study",{})
        if cs and cs.get("title"):
            heading("Case Study", 2)
            heading(cs["title"], 3)
            for lbl, k in [("Challenge","challenge"),("Solution","solution"),("Result","result")]:
                if cs.get(k): para(f"{lbl}: {cs[k]}")

        terms = proposal.get("terms",{})
        if terms:
            doc.add_page_break()
            heading("Terms & Conditions", 2)
            for k, v in terms.items():
                if v:
                    heading(k.replace("_"," ").title(), 3)
                    para(str(v))

        ns = proposal.get("next_steps",[])
        if ns:
            heading("Next Steps", 2)
            for i, step in enumerate(ns,1): bullet(f"{i}. {step}")

        closing = proposal.get("closing_statement","")
        if closing:
            doc.add_paragraph("─"*60)
            para(closing, size=11)

        contacts = [x for x in [branding.get("email",""), branding.get("phone",""), branding.get("website","")] if x]
        if contacts: para(" · ".join(contacts), color=GRAY, size=9)
        para(f"© {datetime.now().year} {agency} · Confidential", color=GRAY, size=8)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        return buf.read()

    except ImportError:
        return b"python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"DOCX error: {e}".encode()
