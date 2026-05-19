"""
parser.py — Extracts text from PDF and DOCX resumes
"""

import io
import os
import re


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
        except Exception as e:
            return f"[Error reading PDF: {e}]"
    except Exception as e:
        return f"[Error reading PDF: {e}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        return f"[Error reading DOCX: {e}]"


def extract_resume_text(uploaded_file) -> str:
    """
    Accept a Streamlit UploadedFile and return extracted text.
    Supports PDF and DOCX.
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        return "[Unsupported file format. Please upload PDF or DOCX.]"
