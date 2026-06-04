"""
Export / rendering layer.

Converts the canonical ``body_json`` newsletter structure into deliverable
formats:

* Markdown  — always available, pure-python
* HTML      — styled, email-friendly inline CSS, always available
* PDF       — via reportlab (graceful message if not installed)
* DOCX      — via python-docx (graceful message if not installed)

Each function returns ``bytes`` (or ``str`` for text formats) so the Streamlit
download buttons can serve them directly.
"""

from __future__ import annotations

import io
from datetime import date


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #
def to_markdown(news: dict) -> str:
    b = news.get("body_json", {})
    lines = [f"# {b.get('title', news.get('title', 'Newsletter'))}", ""]
    if b.get("subject_line"):
        lines += [f"*Subject: {b['subject_line']}*", ""]
    lines += [b.get("introduction", ""), ""]

    if b.get("key_insights"):
        lines += ["## Key Insights", ""]
        for ins in b["key_insights"]:
            if isinstance(ins, dict):
                lines += [f"### {ins.get('heading', '')}", "", ins.get("body", ""), ""]
            else:
                lines += [f"- {ins}", ""]

    if b.get("industry_updates"):
        lines += ["## Industry Updates", ""]
        for up in b["industry_updates"]:
            if isinstance(up, dict):
                head = up.get("headline", "")
                url = up.get("url")
                head_md = f"[{head}]({url})" if url else head
                lines += [f"**{head_md}**", "", up.get("summary", ""), ""]
            else:
                lines += [f"- {up}", ""]

    if b.get("actionable_takeaways"):
        lines += ["## Actionable Takeaways", ""]
        lines += [f"- {t}" for t in b["actionable_takeaways"]]
        lines += [""]

    if b.get("closing"):
        lines += ["## Closing", "", b["closing"], ""]
    if b.get("call_to_action"):
        lines += ["---", "", f"**{b['call_to_action']}**", ""]
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# HTML
# --------------------------------------------------------------------------- #
def to_html(news: dict) -> str:
    b = news.get("body_json", {})

    def section(title, inner):
        return (
            f'<h2 style="font-size:20px;margin:32px 0 12px;color:#111;'
            f'border-bottom:2px solid #6C5CE7;padding-bottom:6px;">{title}</h2>{inner}'
        )

    insights_html = ""
    for ins in b.get("key_insights", []):
        if isinstance(ins, dict):
            insights_html += (
                f'<h3 style="margin:18px 0 6px;color:#2d2d2d;">{ins.get("heading","")}</h3>'
                f'<p style="margin:0 0 12px;line-height:1.6;color:#333;">{ins.get("body","")}</p>'
            )
        else:
            insights_html += f'<p style="line-height:1.6;color:#333;">• {ins}</p>'

    updates_html = ""
    for up in b.get("industry_updates", []):
        if isinstance(up, dict):
            head = up.get("headline", "")
            url = up.get("url")
            head_html = f'<a href="{url}" style="color:#6C5CE7;text-decoration:none;">{head}</a>' if url else head
            updates_html += (
                f'<div style="margin:0 0 16px;padding:14px 16px;background:#f7f7fb;'
                f'border-radius:10px;border-left:3px solid #6C5CE7;">'
                f'<strong style="color:#111;">{head_html}</strong>'
                f'<p style="margin:6px 0 0;line-height:1.55;color:#444;">{up.get("summary","")}</p></div>'
            )

    takeaways_html = "".join(
        f'<li style="margin:0 0 8px;line-height:1.5;color:#333;">{t}</li>'
        for t in b.get("actionable_takeaways", [])
    )

    return f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{b.get('title', 'Newsletter')}</title></head>
<body style="margin:0;padding:0;background:#ececf3;font-family:Georgia,'Times New Roman',serif;">
<div style="max-width:640px;margin:0 auto;background:#ffffff;">
  <div style="background:#6C5CE7;padding:28px 32px;">
    <div style="color:#fff;font-size:13px;letter-spacing:2px;text-transform:uppercase;opacity:.85;">
      {news.get('style','Newsletter')} · {date.today():%B %d, %Y}
    </div>
    <h1 style="color:#fff;font-size:28px;margin:8px 0 0;line-height:1.25;">{b.get('title','Newsletter')}</h1>
  </div>
  <div style="padding:28px 32px;">
    <p style="font-size:17px;line-height:1.65;color:#333;">{b.get('introduction','')}</p>
    {section('Key Insights', insights_html) if insights_html else ''}
    {section('Industry Updates', updates_html) if updates_html else ''}
    {section('Actionable Takeaways', f'<ul style=\"padding-left:20px;\">{takeaways_html}</ul>') if takeaways_html else ''}
    {section('Closing', f'<p style=\"line-height:1.6;color:#333;\">{b.get("closing","")}</p>') if b.get('closing') else ''}
  </div>
  <div style="background:#111;padding:24px 32px;text-align:center;">
    <a href="#" style="display:inline-block;background:#6C5CE7;color:#fff;padding:12px 28px;
       border-radius:30px;text-decoration:none;font-family:Arial,sans-serif;font-size:15px;">
       {b.get('call_to_action','Subscribe')}</a>
  </div>
</div></body></html>"""


# --------------------------------------------------------------------------- #
# PDF (reportlab)
# --------------------------------------------------------------------------- #
def to_pdf(news: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    except ImportError:
        return b"reportlab not installed. Run: pip install reportlab"

    b = news.get("body_json", {})
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    styles = getSampleStyleSheet()
    accent = colors.HexColor("#6C5CE7")
    h1 = ParagraphStyle("h1", parent=styles["Title"], textColor=accent, fontSize=22, spaceAfter=14)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#222"), spaceBefore=16)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=11, leading=16)

    flow = [Paragraph(b.get("title", "Newsletter"), h1)]
    if b.get("subject_line"):
        flow.append(Paragraph(f"<i>{b['subject_line']}</i>", body))
    flow.append(Spacer(1, 10))
    flow.append(Paragraph(b.get("introduction", ""), body))

    if b.get("key_insights"):
        flow.append(Paragraph("Key Insights", h2))
        for ins in b["key_insights"]:
            if isinstance(ins, dict):
                flow.append(Paragraph(f"<b>{ins.get('heading','')}</b>", body))
                flow.append(Paragraph(ins.get("body", ""), body))
            else:
                flow.append(Paragraph(str(ins), body))

    if b.get("industry_updates"):
        flow.append(Paragraph("Industry Updates", h2))
        for up in b["industry_updates"]:
            if isinstance(up, dict):
                flow.append(Paragraph(f"<b>{up.get('headline','')}</b>", body))
                flow.append(Paragraph(up.get("summary", ""), body))

    if b.get("actionable_takeaways"):
        flow.append(Paragraph("Actionable Takeaways", h2))
        flow.append(ListFlowable(
            [ListItem(Paragraph(str(t), body)) for t in b["actionable_takeaways"]],
            bulletType="bullet",
        ))

    if b.get("closing"):
        flow.append(Paragraph("Closing", h2))
        flow.append(Paragraph(b["closing"], body))
    if b.get("call_to_action"):
        flow.append(Spacer(1, 12))
        cta = ParagraphStyle("cta", parent=body, textColor=accent, fontSize=13)
        flow.append(Paragraph(f"<b>{b['call_to_action']}</b>", cta))

    doc.build(flow)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# DOCX (python-docx)
# --------------------------------------------------------------------------- #
def to_docx(news: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        return b"python-docx not installed. Run: pip install python-docx"

    b = news.get("body_json", {})
    doc = Document()
    title = doc.add_heading(b.get("title", "Newsletter"), level=0)
    try:
        title.runs[0].font.color.rgb = RGBColor(0x6C, 0x5C, 0xE7)
    except Exception:
        pass
    if b.get("subject_line"):
        sub = doc.add_paragraph()
        run = sub.add_run(b["subject_line"])
        run.italic = True

    doc.add_paragraph(b.get("introduction", ""))

    if b.get("key_insights"):
        doc.add_heading("Key Insights", level=1)
        for ins in b["key_insights"]:
            if isinstance(ins, dict):
                doc.add_heading(ins.get("heading", ""), level=2)
                doc.add_paragraph(ins.get("body", ""))
            else:
                doc.add_paragraph(str(ins), style="List Bullet")

    if b.get("industry_updates"):
        doc.add_heading("Industry Updates", level=1)
        for up in b["industry_updates"]:
            if isinstance(up, dict):
                p = doc.add_paragraph()
                p.add_run(up.get("headline", "")).bold = True
                doc.add_paragraph(up.get("summary", ""))

    if b.get("actionable_takeaways"):
        doc.add_heading("Actionable Takeaways", level=1)
        for t in b["actionable_takeaways"]:
            doc.add_paragraph(str(t), style="List Bullet")

    if b.get("closing"):
        doc.add_heading("Closing", level=1)
        doc.add_paragraph(b["closing"])
    if b.get("call_to_action"):
        p = doc.add_paragraph()
        p.add_run(b["call_to_action"]).bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


EXPORTERS = {
    "Markdown": (to_markdown, "text/markdown", "md"),
    "HTML": (to_html, "text/html", "html"),
    "PDF": (to_pdf, "application/pdf", "pdf"),
    "DOCX": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
}
